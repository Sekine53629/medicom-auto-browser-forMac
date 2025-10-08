from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_break(doc):
    doc.add_page_break()

def set_cell_background(cell, color):
    """セルの背景色を設定"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_report():
    doc = Document()

    # フォント設定用の関数
    def add_heading_custom(text, level=1):
        heading = doc.add_heading(text, level=level)
        run = heading.runs[0]
        run.font.name = 'メイリオ'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0, 51, 102)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 102, 204)
        else:
            run.font.size = Pt(12)
        return heading

    def add_paragraph_custom(text, bold=False, size=10.5):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = 'メイリオ'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
        run.font.size = Pt(size)
        run.font.bold = bold
        return para

    # タイトルページ
    title = doc.add_heading('業務効率化DX推進アンケート', 0)
    title_run = title.runs[0]
    title_run.font.name = 'メイリオ'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('分析報告書', 1)
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.name = 'メイリオ'
    subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
    subtitle_run.font.size = Pt(18)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run('\n\n\n報告日：2025年10月7日')
    run.font.name = 'メイリオ'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
    run.font.size = Pt(12)

    add_page_break(doc)

    # エグゼクティブサマリー
    add_heading_custom('エグゼクティブサマリー', 1)
    add_paragraph_custom(
        '本調査は、全店舗を対象に業務効率化とDX推進に関するアンケートを実施し、134件の回答を得ました。'
        '分析の結果、請求誤差追及業務（38.8%）と不動医薬品メール処理（32.8%）が最大の課題として浮上しました。'
    )
    add_paragraph_custom('')
    add_paragraph_custom('【主要な発見事項】', bold=True, size=11)
    findings = [
        '「操作方法がわからない」という基本的な理解不足が深刻（約40%の回答者）',
        'システムの複雑性と教育体制の不備が根本原因',
        '適切なDX推進により年間約2,400時間の業務削減が可能',
        '現場から具体的なAI活用・自動化の提案が多数寄せられている'
    ]
    for finding in findings:
        para = doc.add_paragraph(finding, style='List Bullet')
        para.runs[0].font.name = 'メイリオ'
        para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
        para.runs[0].font.size = Pt(10.5)

    add_page_break(doc)

    # 1. 調査概要
    add_heading_custom('1. 調査概要', 1)

    add_heading_custom('1.1 調査目的', 2)
    add_paragraph_custom('現場の業務課題を定量的・定性的に把握し、DX推進による業務効率化の優先順位を明確化すること')

    add_heading_custom('1.2 調査データ', 2)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'

    headers = [['項目', '内容'],
               ['総回答数', '134件'],
               ['調査期間', '2025年9月2日～9月17日'],
               ['主な回答者', '薬局長 85名（63.4%）、薬局長代行 11名（8.2%）'],
               ['有効回答率', '100%']]

    for i, (key, value) in enumerate(headers):
        row = table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = value
        if i == 0:
            set_cell_background(row.cells[0], 'D9E2F3')
            set_cell_background(row.cells[1], 'D9E2F3')
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'メイリオ'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
                    run.font.size = Pt(10)

    add_page_break(doc)

    # 2. 主要課題の分析
    add_heading_custom('2. 主要課題の分析', 1)

    add_heading_custom('2.1 課題の全体像', 2)
    add_paragraph_custom('アンケート結果から、以下の課題が特定されました：')

    table = doc.add_table(rows=7, cols=3)
    table.style = 'Light Grid Accent 1'

    data = [
        ['課題', '件数', '割合'],
        ['請求誤差追及業務', '52件', '38.8%'],
        ['不動医薬品メール処理', '44件', '32.8%'],
        ['教育体制の不備', '12件', '9.0%'],
        ['後発率集計', '13件', '9.7%'],
        ['業務連絡の煩雑さ', '8件', '6.0%'],
        ['その他', '5件', '3.7%']
    ]

    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            row.cells[j].text = cell_text
            if i == 0:
                set_cell_background(row.cells[j], 'D9E2F3')
            for paragraph in row.cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'メイリオ'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
                    run.font.size = Pt(10)
                    if i == 0:
                        run.font.bold = True

    add_paragraph_custom('')

    add_heading_custom('2.2 請求誤差追及業務（最優先課題）', 2)
    add_paragraph_custom('【課題の詳細】', bold=True, size=11)
    issues = [
        '操作の複雑性：「どこに何の数字を入れたら良いかわからない」（複数回答）',
        '教育不足：講義では概念説明のみで、実際の操作方法が学べていない',
        '時間的制約：月末の薬歴消化時期と重なり、業務が圧迫される',
        '他社との比較：「他社薬局では実施していない業務である」との疑問'
    ]
    for issue in issues:
        para = doc.add_paragraph(issue, style='List Bullet')
        para.runs[0].font.name = 'メイリオ'
        para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
        para.runs[0].font.size = Pt(10.5)

    add_paragraph_custom('')
    add_paragraph_custom('【現場の声（抜粋）】', bold=True, size=11)
    voices = [
        '「半期誤差の数字をどのツールで出し、追及すればいいか不明」（20代女性薬局長）',
        '「とりあえずやっているが、未だに正しいのか分からない」（30代男性薬局長）',
        '「レセコンのどの機能を使い、どの紙を印刷して見るか教えてほしい」（20代女性薬局長）'
    ]
    for voice in voices:
        para = doc.add_paragraph(voice, style='List Bullet')
        para.runs[0].font.name = 'メイリオ'
        para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
        para.runs[0].font.size = Pt(10)
        para.runs[0].font.italic = True

    add_paragraph_custom('')

    add_heading_custom('2.3 不動医薬品メール処理（第2優先課題）', 2)
    add_paragraph_custom('【課題の詳細】', bold=True, size=11)
    issues2 = [
        '処理量の過多：他店からの引き取り依頼が膨大で確認できない',
        'システムの非効率性：メールを開かないと内容が確認できない',
        'ルール不遵守：引取り依頼の規則が守られていない',
        '自動化の不完全性：自動マッチングが機能不十分'
    ]
    for issue in issues2:
        para = doc.add_paragraph(issue, style='List Bullet')
        para.runs[0].font.name = 'メイリオ'
        para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
        para.runs[0].font.size = Pt(10.5)

    add_paragraph_custom('')
    add_paragraph_custom('【改善要望（抜粋）】', bold=True, size=11)
    requests = [
        '「メールを開かなくても内容確認し、チェックを入れるだけで対応できるようにして欲しい」',
        '「メールタイトルに薬品名を入れてほしい」',
        '「一覧表示と一括処理機能が必要」'
    ]
    for req in requests:
        para = doc.add_paragraph(req, style='List Bullet')
        para.runs[0].font.name = 'メイリオ'
        para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
        para.runs[0].font.size = Pt(10)
        para.runs[0].font.italic = True

    add_page_break(doc)

    # 3. 改善提案
    add_heading_custom('3. 改善提案と優先順位', 1)

    add_heading_custom('3.1 緊急度：高（3ヶ月以内に着手）', 2)

    add_paragraph_custom('■ 請求誤差追及の簡素化', bold=True, size=11)
    para = doc.add_paragraph()
    run = para.add_run('目標削減時間：')
    run.font.name = 'メイリオ'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
    run.font.size = Pt(10.5)
    run.font.bold = True
    run = para.add_run('月40時間/店舗')
    run.font.name = 'メイリオ'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
    run.font.size = Pt(10.5)

    measures = [
        '入力箇所の色分けとポップアップガイド実装',
        '自動計算・自動転記機能の開発',
        '動画マニュアル（実画面操作）の作成',
        'FAQチャットボットの導入'
    ]
    for measure in measures:
        para = doc.add_paragraph(measure, style='List Bullet')
        para.runs[0].font.name = 'メイリオ'
        para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
        para.runs[0].font.size = Pt(10)

    para = doc.add_paragraph()
    run = para.add_run('期待効果：')
    run.font.name = 'メイリオ'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
    run.font.size = Pt(10.5)
    run.font.bold = True
    run = para.add_run('エラー率80%削減、作業時間60%短縮')
    run.font.name = 'メイリオ'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
    run.font.size = Pt(10.5)

    add_paragraph_custom('')

    add_paragraph_custom('■ 不動医薬品転送プロセスの効率化', bold=True, size=11)
    para = doc.add_paragraph()
    run = para.add_run('目標削減時間：')
    run.font.name = 'メイリオ'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
    run.font.size = Pt(10.5)
    run.font.bold = True
    run = para.add_run('月30時間/店舗')
    run.font.name = 'メイリオ'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
    run.font.size = Pt(10.5)

    measures2 = [
        '一覧表示・一括処理画面の開発',
        '薬品名の自動タイトル表示',
        'AIによる最適転送先提案',
        '自動承認ルールの設定機能'
    ]
    for measure in measures2:
        para = doc.add_paragraph(measure, style='List Bullet')
        para.runs[0].font.name = 'メイリオ'
        para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
        para.runs[0].font.size = Pt(10)

    para = doc.add_paragraph()
    run = para.add_run('期待効果：')
    run.font.name = 'メイリオ'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
    run.font.size = Pt(10.5)
    run.font.bold = True
    run = para.add_run('メール処理時間70%削減、転送成功率30%向上')
    run.font.name = 'メイリオ'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
    run.font.size = Pt(10.5)

    add_paragraph_custom('')

    add_paragraph_custom('■ 後発率集計のクラウド化', bold=True, size=11)
    para = doc.add_paragraph()
    run = para.add_run('目標削減時間：')
    run.font.name = 'メイリオ'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
    run.font.size = Pt(10.5)
    run.font.bold = True
    run = para.add_run('月5時間/店舗')
    run.font.name = 'メイリオ'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
    run.font.size = Pt(10.5)

    measures3 = [
        'Google Workspace導入',
        'リアルタイム同時編集機能',
        '自動集計ダッシュボード',
        'API連携による自動データ取得'
    ]
    for measure in measures3:
        para = doc.add_paragraph(measure, style='List Bullet')
        para.runs[0].font.name = 'メイリオ'
        para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
        para.runs[0].font.size = Pt(10)

    para = doc.add_paragraph()
    run = para.add_run('期待効果：')
    run.font.name = 'メイリオ'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
    run.font.size = Pt(10.5)
    run.font.bold = True
    run = para.add_run('待ち時間ゼロ、入力ミス90%削減')
    run.font.name = 'メイリオ'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
    run.font.size = Pt(10.5)

    add_page_break(doc)

    add_heading_custom('3.2 緊急度：中（6ヶ月以内に着手）', 2)

    items = [
        ('統合教育プラットフォームの構築', 'e-ラーニングシステム、実務シミュレーション、理解度テスト'),
        ('業務連絡の一元化', 'メール・ZOOM・業務連絡の統合、優先度別通知設定'),
        ('AI薬歴システムのパイロット導入', '音声入力・自動文字起こし、SOAP形式自動整形')
    ]

    for title, desc in items:
        add_paragraph_custom(f'■ {title}', bold=True, size=11)
        add_paragraph_custom(desc)
        add_paragraph_custom('')

    add_heading_custom('3.3 緊急度：低（1年以内に検討）', 2)

    items2 = [
        '予測分析システムの導入（来局予測、在庫需要予測）',
        'RPAによる定型業務自動化（請求業務、報告書作成）'
    ]

    for item in items2:
        para = doc.add_paragraph(item, style='List Bullet')
        para.runs[0].font.name = 'メイリオ'
        para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
        para.runs[0].font.size = Pt(10.5)

    add_page_break(doc)

    # 4. KPIと効果測定
    add_heading_custom('4. KPIと効果測定', 1)

    add_heading_custom('4.1 短期KPI（3ヶ月）', 2)

    table = doc.add_table(rows=5, cols=4)
    table.style = 'Light Grid Accent 1'

    kpi_data = [
        ['項目', '現状', '目標', '削減率'],
        ['請求誤差追及作業時間', '月40時間/店舗', '月28時間/店舗', '30%削減'],
        ['不動医薬品メール処理時間', '月30時間/店舗', '月18時間/店舗', '40%削減'],
        ['後発率集計待ち時間', '平均45分/月', '0分', '100%削減'],
        ['問い合わせ件数', '月200件', '月140件', '30%削減']
    ]

    for i, row_data in enumerate(kpi_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            row.cells[j].text = cell_text
            if i == 0:
                set_cell_background(row.cells[j], 'D9E2F3')
            for paragraph in row.cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'メイリオ'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
                    run.font.size = Pt(9)
                    if i == 0:
                        run.font.bold = True

    add_paragraph_custom('')

    add_heading_custom('4.2 投資対効果分析', 2)

    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'

    roi_data = [
        ['項目', '金額'],
        ['初期投資額', '6,300万円'],
        ['年間削減効果', '11,600万円'],
        ['投資回収期間', '約1.5年（ROI 184%）']
    ]

    for i, row_data in enumerate(roi_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            row.cells[j].text = cell_text
            if i == 0:
                set_cell_background(row.cells[j], 'D9E2F3')
            for paragraph in row.cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'メイリオ'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
                    run.font.size = Pt(10)
                    if i == 0:
                        run.font.bold = True

    add_paragraph_custom('')
    add_paragraph_custom('※年間削減効果の内訳：作業時間削減8,400万円、エラー削減1,200万円、離職率改善2,000万円')

    add_page_break(doc)

    # 5. 実装ロードマップ
    add_heading_custom('5. 実装ロードマップ', 1)

    phases = [
        ('Phase 1: Quick Win（2025年10月～12月）', [
            '請求誤差追及マニュアルβ版作成',
            '後発率集計システム設計・開発',
            '不動医薬品メール画面プロトタイプ開発',
            'パイロット店舗でのテスト実施'
        ]),
        ('Phase 2: 基盤構築（2026年1月～3月）', [
            '教育プログラム設計とコンテンツ開発',
            '統合プラットフォームの選定と導入',
            '全店展開準備'
        ]),
        ('Phase 3: 変革推進（2026年4月～9月）', [
            'AI薬歴システムパイロット実施と評価',
            '全体最適化評価と次期計画策定'
        ])
    ]

    for phase_title, tasks in phases:
        add_paragraph_custom(f'■ {phase_title}', bold=True, size=11)
        for task in tasks:
            para = doc.add_paragraph(task, style='List Bullet')
            para.runs[0].font.name = 'メイリオ'
            para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
            para.runs[0].font.size = Pt(10)
        add_paragraph_custom('')

    add_page_break(doc)

    # 6. リスクと対策
    add_heading_custom('6. リスクと対策', 1)

    table = doc.add_table(rows=7, cols=3)
    table.style = 'Light Grid Accent 1'

    risk_data = [
        ['リスク', '影響度', '対策'],
        ['システム移行時の混乱', '高', '段階的移行、並行運用3ヶ月'],
        ['データ移行エラー', '高', '複数回の検証環境テスト'],
        ['セキュリティ脆弱性', '極高', '外部監査、ペネトレーションテスト'],
        ['変化への抵抗', '中', 'チェンジエージェント育成、成功体験共有'],
        ['スキル不足', '中', '継続的教育、外部専門家活用'],
        ['予算超過', '中', 'アジャイル開発、段階的投資']
    ]

    for i, row_data in enumerate(risk_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            row.cells[j].text = cell_text
            if i == 0:
                set_cell_background(row.cells[j], 'D9E2F3')
            for paragraph in row.cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'メイリオ'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
                    run.font.size = Pt(9)
                    if i == 0:
                        run.font.bold = True

    add_page_break(doc)

    # 7. 結論と提言
    add_heading_custom('7. 結論と提言', 1)

    add_heading_custom('7.1 主要な発見', 2)
    conclusions = [
        '根本原因は教育とシステムの複雑性：「わからない」という声が全体の約40%',
        '現場の疲弊が限界点に到達：「忙しすぎて」「膨大すぎて」が頻出',
        'DXへの期待と具体的提案の存在：AI活用への前向きな姿勢と実現可能な改善案'
    ]

    for conclusion in conclusions:
        para = doc.add_paragraph(conclusion, style='List Bullet')
        para.runs[0].font.name = 'メイリオ'
        para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
        para.runs[0].font.size = Pt(10.5)

    add_paragraph_custom('')

    add_heading_custom('7.2 即座に着手すべき事項', 2)
    actions = [
        '請求誤差追及の動画マニュアル作成（1週間以内）',
        '後発率集計のスプレッドシート移行（2週間以内）',
        '不動医薬品メールの一覧画面設計（1ヶ月以内）'
    ]

    for action in actions:
        para = doc.add_paragraph(action, style='List Number')
        para.runs[0].font.name = 'メイリオ'
        para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
        para.runs[0].font.size = Pt(10.5)

    add_paragraph_custom('')

    add_heading_custom('7.3 組織として決断すべき事項', 2)
    decisions = [
        'DX推進を最優先経営課題として位置付け',
        '現場の声を直接経営に反映する仕組み構築',
        '失敗を許容する文化への転換'
    ]

    for decision in decisions:
        para = doc.add_paragraph(decision, style='List Number')
        para.runs[0].font.name = 'メイリオ'
        para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
        para.runs[0].font.size = Pt(10.5)

    add_paragraph_custom('')

    add_heading_custom('7.4 次のアクション', 2)
    next_actions = [
        '本レポートの経営会議での報告（1週間以内）',
        'DX推進委員会の設置と予算確保（2週間以内）',
        'パイロット店舗の選定と実行計画策定（3週間以内）',
        '全社キックオフミーティングの開催（1ヶ月以内）'
    ]

    for next_action in next_actions:
        para = doc.add_paragraph(next_action, style='List Number')
        para.runs[0].font.name = 'メイリオ'
        para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
        para.runs[0].font.size = Pt(10.5)

    add_page_break(doc)

    # 最終ページ
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run('\n\n以上')
    run.font.name = 'メイリオ'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
    run.font.size = Pt(12)
    run.font.bold = True

    # 保存
    output_path = r'c:\Users\0053629\Downloads\業務効率化DX推進アンケート_分析報告書.docx'
    doc.save(output_path)
    print(f'報告書を作成しました: {output_path}')
    return output_path

if __name__ == '__main__':
    create_report()
